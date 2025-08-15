# app/services/docx_processor.py

import os
import json
import re
import hashlib
from cryptography.fernet import Fernet
from docx import Document
from app.services.pii_main import extract_all_pii

def mask_docx_sensitive_text(docx_path: str, key_path: str = None, enabled_pii_categories=None):
    if key_path is None:
        key_path = docx_path.replace(".docx", ".key")

    # Load or generate a key
    if os.path.exists(key_path):
        with open(key_path, "rb") as f:
            key = f.read()
    else:
        key = Fernet.generate_key()
        with open(key_path, "wb") as f:
            f.write(key)

    fernet = Fernet(key)
    document = Document(docx_path)

    # Extract all text from the document first
    full_text = ""
    for para in document.paragraphs:
        full_text += para.text + "\n"

    # Extract all PII at once to avoid duplicates
    all_pii_list = extract_all_pii(full_text, enabled_pii_categories)

    # Create unique PII mapping with hash-based tags
    unique_pii = {}
    masked_pii = []

    for label, value in all_pii_list:
        if value not in unique_pii:
            try:
                encrypted = fernet.encrypt(value.encode()).decode()
                # Use hash for unique tags like in text processor
                import hashlib
                value_hash = hashlib.md5(value.encode()).hexdigest()[:8]
                unique_tag = f"[ENC:{label}_{value_hash}]"

                unique_pii[value] = {"encrypted": encrypted, "tag": unique_tag}
                masked_pii.append({
                    "original": value,
                    "encrypted": encrypted,
                    "label": label,
                    "masked": unique_tag
                })
            except Exception as e:
                print(f"[ERROR] Failed to encrypt '{value}': {e}")

    # Sort by length for safe replacement
    sorted_pii_items = sorted(unique_pii.items(), key=lambda x: len(x[0]), reverse=True)

    # Apply safe replacement to each paragraph
    for para in document.paragraphs:
        if para.text.strip():  # Only process non-empty paragraphs
            masked_text = para.text

            for pii_value, pii_info in sorted_pii_items:
                # Use the same safe replacement logic as text processor
                parts = re.split(r'(\[ENC:[^\]]+\])', masked_text)

                for i in range(len(parts)):
                    if i % 2 == 0 and not parts[i].startswith('[ENC:'):
                        escaped_pii = re.escape(pii_value)
                        if len(pii_value) == 1 and pii_value.isdigit():
                            pattern = r'\b' + escaped_pii + r'\b'
                        else:
                            pattern = escaped_pii
                        parts[i] = re.sub(pattern, pii_info["tag"], parts[i])

                masked_text = ''.join(parts)

            para.text = masked_text

    # Output File
    masked_path = docx_path.replace(".docx", ".masked.docx")
    document.save(masked_path)

    json_path = docx_path.replace(".docx", ".masked.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(masked_pii, f, ensure_ascii=False, indent=2)

    return masked_path, json_path, key_path

def run_docx_processing(docx_path: str, enabled_pii_categories=None):
    try:
        key_path = docx_path.replace(".docx", ".key")
        masked_docx, json_path, key_file = mask_docx_sensitive_text(
            docx_path, key_path, enabled_pii_categories
        )
        return {
            "status": "success",
            "masked_docx": masked_docx,
            "json_output": json_path,
            "key_file": key_file
        }
    except Exception as e:
        return {
            "status": "error",
            "message": str(e)
        }
